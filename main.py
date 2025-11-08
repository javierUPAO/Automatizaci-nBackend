from fastapi import FastAPI, UploadFile,Form,Request
from fastapi.responses import FileResponse
from fastapi.responses import StreamingResponse
from fastapi.responses import JSONResponse
import base64
import json
from docx import Document
import tempfile
import uvicorn
import re
import logging
import time
from io import BytesIO
import threading, uuid, requests
from fastapi.responses import StreamingResponse
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor, Inches
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT


app = FastAPI() 

logging.basicConfig(level=logging.INFO, format='%(asctime)s [%(levelname)s] %(message)s')

#------------------------------------------------------------------------------------#   
@app.post("/b64ToDocxN8n")
async def process_doc_base64(file_base64: str = Form(...)):
    try:
        texto_bytes = file_base64.encode("utf-8")  # convertir a bytes
        texto_b64 = base64.b64encode(texto_bytes).decode("utf-8")  # codificar y volver a string
        # Decodificar texto base64
        decoded_text = base64.b64decode(texto_b64).decode("utf-8")

        # Crear documento
        doc = Document()

        # Márgenes más amplios
        for section in doc.sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)

        # Título principal
        title = doc.add_heading("Plan de Curso", level=1)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title.runs[0].font.size = Pt(20)
        title.runs[0].font.color.rgb = RGBColor(0, 102, 204)  # azul elegante

        # Separar por líneas
        lines = decoded_text.splitlines()

        for line in lines:
            line = line.strip()
            if not line:
                continue  # saltar líneas vacías

            # === Título principal ===
            if line.lower().startswith("sílabus del curso"):
                heading = doc.add_heading(line, level=1)
                run = heading.runs[0]
                run.font.color.rgb = RGBColor(0, 51, 102)
                run.font.size = Pt(16)
                run.font.bold = True
                doc.add_paragraph("")  # Espacio después

            # === Información General ===
            elif line.lower().startswith("información general"):
                heading = doc.add_heading(line, level=2)
                run = heading.runs[0]
                run.font.color.rgb = RGBColor(0, 51, 102)
                run.font.size = Pt(14)
                run.font.bold = True

            # === Elementos con sangría dentro de Información General ===
            elif any(line.lower().startswith(x) for x in ["duración", "modalidad"]):
                p = doc.add_paragraph(style='List Bullet')
                run = p.add_run(line)
                run.font.size = Pt(11)
                p.paragraph_format.left_indent = Inches(0.5)
                p.paragraph_format.space_after = Pt(3)

            # === Plan de Estudios ===
            elif line.lower().startswith("plan de estudios"):
                heading = doc.add_heading(line, level=2)
                run = heading.runs[0]
                run.font.color.rgb = RGBColor(0, 51, 102)
                run.font.size = Pt(14)
                run.font.bold = True

            elif re.match(r"^módulo\s+\d+", line.lower()):
                heading = doc.add_heading(line, level=3)
                run = heading.runs[0]
                run.font.color.rgb = RGBColor(0, 102, 204)
                run.font.size = Pt(13)
                run.font.bold = True
                keep_with_next(heading)
                doc.add_paragraph("")

            # === Lecciones ===
            elif re.match(r"^lección\s+\d+\.\d+", line.lower()):
                p = doc.add_paragraph(style='List Bullet')
                run = p.add_run(line)
                run.font.color.rgb = RGBColor(0, 0, 128)
                run.font.bold = True
                run.font.size = Pt(12)
                p.paragraph_format.left_indent = Inches(0.5)
                p.paragraph_format.space_after = Pt(2)
                keep_with_next(p)

            # === Objetivos dentro de Lecciones ===
            elif line.lower().startswith("objetivo:"):
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Inches(1)
                p.paragraph_format.line_spacing = 1.3
                p.paragraph_format.space_after = Pt(6)
                p.add_run("Objetivo: ").bold = True
                p.add_run(line.replace("Objetivo:", "").strip())
                keep_lines_together(p)
                doc.add_paragraph("")

            # === Texto normal (descripciones u observaciones) ===
            else:
                p = doc.add_paragraph(line)
                p.paragraph_format.left_indent = Inches(0.3)
                p.paragraph_format.space_after = Pt(4)
                p.paragraph_format.line_spacing = 1.2

        # Guardar documento en memoria
        output_io = BytesIO()
        doc.save(output_io)
        output_io.seek(0)

        # Retornar DOCX descargable
        return StreamingResponse(
            output_io,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": "attachment; filename=plan_curso.docx"}
        )

    except Exception as e:
        return {"error": str(e)}

def keep_lines_together(paragraph):
    pPr = paragraph._element.get_or_add_pPr()
    keepLines = OxmlElement('w:keepLines')
    pPr.append(keepLines)
#------------------------------------------------------------------------------------#  
def keep_with_next(paragraph):
    pPr = paragraph._element.get_or_add_pPr()
    keepNext = OxmlElement('w:keepNext')
    pPr.append(keepNext)

@app.post("/generate-doc")
async def generate_doc(data: dict):
    doc = Document()
    p = doc.add_heading(level=1)  # nivel 1
    run = p.add_run(data["curso"])
    run.font.size = Pt(18)
    run.bold = True
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    doc.add_paragraph(f"Duración: {data['duracion']}")
    doc.add_paragraph(f"Modalidad: {data['modalidad']}")
    
    for modulo in data.get("modulos", []):
        # --- Módulo ---
        p_mod = doc.add_paragraph()
        run_mod = p_mod.add_run(f"Módulo: {modulo['nombre']}")
        run_mod.font.size = Pt(14)
        run_mod.bold = True
        run_mod.font.color.rgb = RGBColor(79, 129, 189) 
        p_mod.paragraph_format.left_indent = Pt(0)  # nivel 0


        for leccion in modulo.get("lecciones", []):
            # --- Lección ---
            p_lec = doc.add_paragraph()
            run_lec = p_lec.add_run(f"Lección: {leccion['titulo']}")
            run_lec.font.size = Pt(13)
            run_lec.bold = True
            run_lec.font.color.rgb = RGBColor(79, 129, 189)
            p_lec.paragraph_format.left_indent = Pt(18)  # sangría jerárquica


            for obj in leccion.get("objetivos", []):
                p_obj = doc.add_paragraph(f"{obj}")
                p_obj.paragraph_format.left_indent = Pt(36)  # sangría
                p_obj.paragraph_format.space_after = Pt(3)
                run_obj = p_obj.runs[0]
                run_obj.font.size = Pt(11)
            
            contenido = leccion.get("contenido", {})
            if contenido:
                agenda = contenido.get("agenda", [])
                if agenda:
                    p_ag = doc.add_paragraph("Agenda:")
                    p_ag.paragraph_format.left_indent = Pt(36)
                    run_ag = p_ag.runs[0]
                    run_ag.font.color.rgb = RGBColor(79, 129, 189)
                    run_ag.font.size = Pt(13)
                    run_ag.bold = True
                    keep_with_next(p_ag)
                # Agenda
                    for tema in agenda:
                        p_tema = doc.add_paragraph(tema, style='List Bullet')
                        p_tema.paragraph_format.left_indent = Pt(70)
                        p_tema.paragraph_format.space_after = Pt(2)
                        run_tema = p_tema.runs[0]
                        run_tema.font.size = Pt(11)
                
                # Desarrollo de temas
                for tema in contenido.get("desarrollo_temas", []):
                    p_titulo = doc.add_paragraph(tema.get("titulo_tema", ""))
                    p_titulo.paragraph_format.left_indent = Pt(36)
                    run_titulo = p_titulo.runs[0]
                    run_titulo.font.color.rgb = RGBColor(79, 129, 189)
                    run_titulo.font.size = Pt(13)
                    run_titulo.bold = True
                    keep_with_next(p_titulo)

                    p_cont = doc.add_paragraph(tema.get("contenido_tema", ""))
                    p_cont.paragraph_format.left_indent = Pt(54)
                    run_cont = p_cont.runs[0]
                    run_cont.font.size = Pt(11)
                
                # Actividad práctica
                actividad = contenido.get("actividad_practica")
                if actividad:
                    p_act = doc.add_paragraph(f"Actividad: {actividad.get('titulo_actividad', '')}")
                    p_act.paragraph_format.left_indent = Pt(36)
                    run_act = p_act.runs[0]
                    run_act.font.color.rgb = RGBColor(79, 129, 189)
                    run_act.font.size = Pt(13)
                    run_act.bold = True
                    keep_with_next(p_act)

                    p_desc = doc.add_paragraph(actividad.get("descripcion_actividad", ""))
                    p_desc.paragraph_format.left_indent = Pt(54)
                    run_desc = p_desc.runs[0]
                    run_desc.font.size = Pt(11)
                
                # Referencias
                referencias = contenido.get("referencias", [])
                if referencias:
                    p_ref = doc.add_paragraph("Referencias")
                    p_ref.paragraph_format.left_indent = Pt(36)
                    run_ref = p_ref.runs[0]
                    run_ref.font.color.rgb = RGBColor(79, 129, 189)
                    run_ref.font.size = Pt(13)
                    run_ref.bold = True
                    keep_with_next(p_ref)

                    for ref in referencias:
                        p = doc.add_paragraph(ref)
                        # Configurar sangría francesa
                        p.paragraph_format.left_indent = Pt(73)          # margen izquierdo
                        p.paragraph_format.first_line_indent = Pt(-18) # sangría francesa de 1/2 pulgada aprox.
                        p.paragraph_format.space_after = Pt(6)         # espacio después del párrafo

    # Guardar en memoria
    file_stream = BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)

    # Devolver archivo como Word
    filename = "".join(c if c.isalnum() or c in "_-" else "_" for c in data["curso"])
    return StreamingResponse(
        file_stream,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename={filename}.docx"}
    )

##return FileResponse(output_path, filename="resultado.docx")

@app.post("/textExtractor")
async def process_doc(file: UploadFile):
    logging.info(f"Archivo recibido: {file.filename}")
    
     # Guardar archivo temporal
    with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
        tmp.write(await file.read())
        tmp_path = tmp.name
    logging.info(f"Archivo guardado temporalmente en {tmp_path}")

    # Extraer texto
    doc = Document(tmp_path)
    input_text = "\n\n".join([p.text.strip() for p in doc.paragraphs if p.text.strip()])
    logging.info(f"Texto extraído: {len(input_text)} caracteres")

    # Devolver JSON con el archivo codificado
    return JSONResponse(content={
        "message": "Texto extraido correctamente",
        "Texto": input_text
    })

## ---------------##
import asyncio
# Lista de clientes conectados
clients = []

# Función que genera eventos SSE para cada cliente
async def event_generator(client_queue: asyncio.Queue):
    try:
        while True:
            # Espera hasta que haya datos para enviar
            data = await client_queue.get()
            yield f"data: {json.dumps(data)}\n\n"
    except asyncio.CancelledError:
        # Si se desconecta el cliente
        pass

# Endpoint SSE que Angular va a escuchar
@app.get("/sse")
async def sse():
    async def event_generator():
        queue = asyncio.Queue()
        clients.append(queue)
        try:
            while True:
                data = await queue.get()
                yield f"data: {json.dumps(data)}\n\n"
        except asyncio.CancelledError:
            pass
        finally:
            clients.remove(queue)
    return StreamingResponse(event_generator(), media_type="text/event-stream")

# Función para enviar datos a todos los clientes conectados
def send_event_to_clients(data: dict):
    for client_queue in clients:
        client_queue.put_nowait(data)

# Endpoint de callback de n8n
@app.post("/callback-n8n")
async def callback_n8n(payload: dict):
    send_event_to_clients(payload)
    return {"status": "ok"}


#-----Cors policy------#
from fastapi.middleware.cors import CORSMiddleware

app.add_middleware(
    CORSMiddleware,
    allow_origins=[
        "http://localhost:4200",
        "https://superauto-75513.web.app"
        ],  
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

if __name__ == "__main__":
    port = int(os.environ.get("PORT", 8000))  # Railway asigna el puerto
    uvicorn.run(app, host="0.0.0.0", port=port)